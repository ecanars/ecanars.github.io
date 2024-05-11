from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.http import require_POST
from reportlab.pdfgen import canvas
from docx import Document
import io
from.models import *
    
courses = [
    "CS101", 
    "CS102",
    "CS201",
    "CS202",
    "CS319",
    "CS315",
]

def create_question_view(request):
    template_name = 'question-form.html'
    context={
        'courses' : courses
    }
    if request.method == 'POST':
        
        prompt = request.POST["prompt"]
        answer = request.POST["answer"]
        category = request.POST["category"]
        difficulty = request.POST["difficulty"]
        
        question = Question.create(prompt=prompt, answer=answer, category=category, difficulty=difficulty)
        question.save()
    return render(request, template_name=template_name, context=context )

def question_list_view(request):
    template_name = 'question-list.html'
    questions = Question.objects.all()
    context = {
        'questions': questions
    }
    return render(request, template_name=template_name, context=context)

def question_detail_view(request, pk):
    template_name = 'question-detail.html'
    question = Question.objects.get(pk=pk)
    context = {
        'question': question
    }
    if request.method == 'POST':
        question.delete()
        return redirect("/question-list")
    return render(request, template_name=template_name, context=context)

def main(request):
    return render(request, 'base.html')

def new_assessment_view_1(request):
    template_name = 'new-assessment_1.html'
    context = {
        'courses': courses
    }
    if request.method == 'POST':
        name = request.POST["name"]
        course = request.POST["course"]
        assessment = Assessment.create(name, course)
        assessment.save()
        next_url = "/assessment-add-question/" + str(assessment.pk)
        return redirect(next_url)
    
    return render(request, template_name=template_name, context=context)

def assessment_add_question_view(request, pk):
    template_name = 'assessment-add-question.html'
    assessment = Assessment.objects.get(pk=pk)
    assessment_questions = assessment.questions.all()
    question_count = len(assessment_questions)
    questions_add = Question.objects.filter(category=assessment.course)
    context = {
        'assessment': assessment,
        'assessment_questions': assessment_questions,
        'question_count': question_count,
        'questions_add': questions_add,
    }
    if request.method == 'POST':
        if 'question-add' in request.POST:
            question_id = request.POST["question-add"]
            question = Question.objects.get(id=question_id)
            assessment.questions.add(question)
            assessment.save()
        elif 'remove-question' in request.POST:
            question_id = request.POST["remove-question"]
            question = Question.objects.get(id=question_id)
            assessment.questions.remove(question)
            assessment.save()
        next_url = "/assessment-add-question/" + str(assessment.pk)
        return redirect(next_url)
    return render(request, template_name=template_name, context=context)

def assessment_list_view(request):
    template_name = 'assessment-list.html'
    assessments = Assessment.objects.all()
    context = {
        'assessments': assessments
    }
    return render(request, template_name=template_name, context=context)

def assessment_detail_view(request, pk):
    template_name = 'assessment-detail.html'
    assessment = Assessment.objects.get(pk=pk)
    questions = assessment.questions.all()
    question_count = len(questions)
    context = {
        'assessment': assessment,
        'questions': questions,
        'question-count': question_count,
    }
    if request.method == 'POST':
        assessment.delete()
        return redirect("/assessment-list")
    return render(request, template_name=template_name, context=context)

@require_POST
def export_assessment(request, pk):
    assessment = get_object_or_404(Assessment, pk=pk)
    export_format = request.POST.get('format')

    if export_format == 'pdf':
        response = export_to_pdf(request,pk)
    elif export_format == 'docx':
        response = export_to_docx(request,pk)
    else:
        return HttpResponse("Invalid format", status=400)

    return response

def export_to_pdf(request, pk):
    # Fetch the assessment by primary key or return a 404 error if not found
    try:
        assessment = Assessment.objects.get(pk=pk)
    except Assessment.DoesNotExist:
        return HttpResponse("The requested assessment does not exist.", status=404)

    # Prepare the PDF response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{assessment.name}.pdf"'

    # Create a PDF canvas and start writing data
    p = canvas.Canvas(response)
    title_y = 800
    question_y = 750
    p.drawString(100, title_y, f"Assessment: {assessment.name}")
    p.drawString(100, title_y - 30, f"Course: {assessment.course}")

    # Output questions
    for idx, question in enumerate(assessment.questions.all(), 1):
        p.drawString(100, question_y - (40 * idx), f"Q{idx}: {question.prompt}")
        p.drawString(100, question_y - (40 * idx) - 20, f"Answer: {question.answer}")

    p.showPage()
    p.save()
    return response


def export_to_docx(request, pk):
    try:
        assessment = Assessment.objects.get(pk=pk)
        document = Document()
        
        # Title and course of the assessment
        document.add_heading('Assessment Details', level=1)
        document.add_paragraph(f"Name: {assessment.name}")
        document.add_paragraph(f"Course: {assessment.course}")

        # Adding a section for questions
        document.add_heading('Questions', level=2)
        for idx, question in enumerate(assessment.questions.all(), 1):
            document.add_paragraph(f"Q{idx}: {question.prompt}")
            document.add_paragraph(f"Answer: {question.answer}")

        # Prepare response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{assessment.name}.docx"'
        document.save(response)

        return response
    except Assessment.DoesNotExist:
        return HttpResponse("The requested assessment does not exist.", status=404)


